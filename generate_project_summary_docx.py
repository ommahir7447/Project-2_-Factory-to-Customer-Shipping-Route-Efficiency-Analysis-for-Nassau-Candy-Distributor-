"""
Generate In-Depth Project Summary -- Google Docs-compatible .docx
=================================================================
Nassau Candy Distributor: Factory-to-Customer Shipping Route Efficiency Analysis

Run:  python generate_project_summary_docx.py
Output: Project_Summary_Nassau_Candy.docx  (same directory)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# -----------------------------------------------
# COLOUR PALETTE
# -----------------------------------------------
PRIMARY       = RGBColor(0x0A, 0x26, 0x47)
ACCENT        = RGBColor(0x00, 0x89, 0x7B)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
TEXT          = RGBColor(0x21, 0x21, 0x21)
TEXT_SEC      = RGBColor(0x54, 0x6E, 0x7A)
DANGER        = RGBColor(0xD3, 0x2F, 0x2F)
WARNING       = RGBColor(0xF5, 0x7C, 0x00)
SUCCESS       = RGBColor(0x2E, 0x7D, 0x32)
BLUE          = RGBColor(0x15, 0x65, 0xC0)
HDR_BG        = "0A2647"
ALT_ROW       = "F5F7FA"
ACCENT_BG     = "E0F2F1"
BORDER        = "BDBDBD"
FONT          = "Calibri"


# -----------------------------------------------
# HELPER UTILITIES
# -----------------------------------------------
def shade_cell(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    )

def border_cell(cell, color=BORDER):
    bs = {"sz": "4", "val": "single", "color": color}
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge in ("top", "bottom", "left", "right"):
        tcBorders.append(parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{bs["val"]}" '
            f'w:sz="{bs["sz"]}" w:space="0" w:color="{bs["color"]}"/>'
        ))
    cell._tc.get_or_add_tcPr().append(tcBorders)

def styled_para(doc, text, size=11, bold=False, italic=False,
                color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT,
                before=0, after=6, indent=None, bg=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if bg:
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'))
    run = p.add_run(text)
    run.font.size     = Pt(size)
    run.font.bold     = bold
    run.font.italic   = italic
    run.font.color.rgb = color
    run.font.name     = FONT
    return p

def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(8)
    run_n = p.add_run(f"{number}.  ")
    run_n.font.size = Pt(15); run_n.font.bold = True
    run_n.font.color.rgb = ACCENT; run_n.font.name = FONT
    run_t = p.add_run(title)
    run_t.font.size = Pt(15); run_t.font.bold = True
    run_t.font.color.rgb = PRIMARY; run_t.font.name = FONT
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="00897B"/>'
        f'</w:pBdr>'
    ))
    return p

def sub_heading(doc, title, level=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(title)
    run.font.size = Pt(12); run.font.bold = True
    run.font.color.rgb = PRIMARY; run.font.name = FONT
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        c = t.cell(0, i); c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.font.size = Pt(9); run.font.bold = True
        run.font.color.rgb = WHITE; run.font.name = FONT
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(c, HDR_BG); border_cell(c, HDR_BG)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.cell(ri + 1, ci); c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10); run.font.name = FONT
            run.font.color.rgb = TEXT
            if ci == 0:
                run.font.bold = True; run.font.color.rgb = PRIMARY
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                shade_cell(c, ALT_ROW)
            border_cell(c)
    doc.add_paragraph()  # spacer
    return t

def bullet_block(doc, icon, title, body, color=TEXT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{icon}  "); r1.font.size = Pt(11); r1.font.name = FONT
    r1.font.color.rgb = color
    r2 = p.add_run(title)
    r2.font.size = Pt(11); r2.font.bold = True; r2.font.name = FONT
    r2.font.color.rgb = PRIMARY
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(8)
    p2.paragraph_format.left_indent  = Cm(1.0)
    r3 = p2.add_run(body)
    r3.font.size = Pt(10); r3.font.name = FONT; r3.font.color.rgb = TEXT_SEC

def divider(doc):
    r = doc.add_paragraph().add_run("-" * 90)
    r.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
    r.font.size = Pt(6)


# ==================================================
# MAIN DOCUMENT
# ==================================================
def generate():
    doc = Document()

    # Page setup (A4)
    sec = doc.sections[0]
    sec.page_width   = Cm(21.0)
    sec.page_height  = Cm(29.7)
    sec.top_margin   = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin  = Cm(2.0)
    sec.right_margin = Cm(2.0)

    # Default style
    ns = doc.styles["Normal"]
    ns.font.name = FONT; ns.font.size = Pt(11); ns.font.color.rgb = TEXT

    # =========================================================
    # COVER / TITLE BLOCK
    # =========================================================
    styled_para(doc, "PROJECT SUMMARY REPORT", 9, bold=True, color=ACCENT, before=4, after=4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Factory-to-Customer Shipping Route\nEfficiency Analysis")
    r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = PRIMARY; r.font.name = FONT

    styled_para(doc, "Nassau Candy Distributor  --  Comprehensive Logistics Performance Review", 13,
                color=TEXT_SEC, before=0, after=10)

    # Metadata
    meta_items = [
        ("Project Type:", "Data Analytics & Supply Chain Optimization"),
        ("Date:", "July 2026"),
        ("Analyst:", "Om Ahir"),
        ("Prepared for:", "Government & Regulatory Stakeholders"),
        ("Tools Used:", "Python (Pandas, Plotly), Streamlit, Power BI, MS Excel"),
    ]
    for label, value in meta_items:
        pm = doc.add_paragraph()
        pm.paragraph_format.space_before = Pt(0)
        pm.paragraph_format.space_after  = Pt(2)
        rl = pm.add_run(label + " ")
        rl.font.size = Pt(10); rl.font.bold = True; rl.font.color.rgb = PRIMARY; rl.font.name = FONT
        rv = pm.add_run(value)
        rv.font.size = Pt(10); rv.font.color.rgb = TEXT_SEC; rv.font.name = FONT

    # Confidential
    styled_para(doc, "CONFIDENTIAL -- For authorized government stakeholder review only", 8,
                bold=True, color=RGBColor(0xF5, 0x7F, 0x17),
                align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=16, bg="FFF8E1")

    divider(doc)

    # =========================================================
    # TABLE OF CONTENTS
    # =========================================================
    styled_para(doc, "TABLE OF CONTENTS", 14, bold=True, color=PRIMARY, before=8, after=10)
    toc = [
        "1.  Project Overview & Objectives",
        "2.  Dataset Description & Data Engineering",
        "3.  Methodology & Analytical Framework",
        "4.  Key Performance Indicators (KPIs)",
        "5.  Critical Findings & Insights",
        "6.  Geographic Impact Analysis",
        "7.  Ship Mode Performance Analysis",
        "8.  Product & Division Analysis",
        "9.  Temporal & Seasonal Patterns",
        "10. Financial & Profitability Analysis",
        "11. Risk Assessment",
        "12. Strategic Recommendations",
        "13. Implementation Roadmap",
        "14. Technical Architecture & Deliverables",
        "15. Conclusion",
    ]
    for item in toc:
        styled_para(doc, item, 10, color=PRIMARY, indent=0.5, before=1, after=1)

    doc.add_page_break()

    # =========================================================
    # 1. PROJECT OVERVIEW & OBJECTIVES
    # =========================================================
    section_heading(doc, 1, "Project Overview & Objectives")

    sub_heading(doc, "1.1  Background")
    styled_para(doc,
        "Nassau Candy Distributor is a U.S.-based confectionery manufacturer and wholesale distributor "
        "that ships products from its northeastern factory to customers across all 50 U.S. states and "
        "territories. As the company scales, understanding the efficiency of its factory-to-customer "
        "shipping routes becomes critical for maintaining competitive service levels, controlling "
        "logistics costs, and ensuring equitable delivery experiences across geographic regions.",
        10.5, color=TEXT, before=4, after=8)

    sub_heading(doc, "1.2  Problem Statement")
    styled_para(doc,
        "The company faces significant challenges in its shipping operations: extreme variability in "
        "delivery times across regions, a high overall delay rate of approximately 80%, reliance on a "
        "single factory distribution model, and limited visibility into route-level performance. These "
        "inefficiencies result in customer dissatisfaction, regional delivery inequality, and suboptimal "
        "logistics spend.",
        10.5, color=TEXT, before=4, after=8)

    sub_heading(doc, "1.3  Project Objectives")
    objectives = [
        ("Data Engineering:", "Transform raw transactional data (18 columns) into an analysis-ready "
         "dataset (43 engineered features) with computed shipping metrics, route classifications, "
         "efficiency scores, and financial ratios."),
        ("Route Efficiency Analysis:", "Evaluate and rank factory-to-customer shipping routes by "
         "lead time, on-time performance, and a composite efficiency score across all U.S. regions."),
        ("Geographic Disparity Identification:", "Quantify regional differences in delivery performance "
         "and identify underserved corridors requiring logistics intervention."),
        ("Ship Mode Optimization:", "Analyze the cost-speed tradeoff across four shipping tiers "
         "(Same Day, First Class, Second Class, Standard Class) to identify upgrade opportunities."),
        ("Interactive Dashboard:", "Build a production-grade Streamlit dashboard with 6 analytical tabs, "
         "interactive filters, and real-time KPI visualization for stakeholder decision-making."),
        ("Actionable Recommendations:", "Deliver evidence-based strategic recommendations with a phased "
         "implementation roadmap to improve delivery performance by 15-20%."),
    ]
    for title, body in objectives:
        bullet_block(doc, ">>", title, body, ACCENT)

    doc.add_page_break()

    # =========================================================
    # 2. DATASET DESCRIPTION & DATA ENGINEERING
    # =========================================================
    section_heading(doc, 2, "Dataset Description & Data Engineering")

    sub_heading(doc, "2.1  Raw Dataset Overview")
    styled_para(doc,
        "The source dataset (Nassau Candy Distributor.csv) contains 10,194 transactional records "
        "with 18 original columns capturing order details, customer information, geographic data, "
        "product details, and financial metrics.",
        10.5, color=TEXT, before=4, after=8)

    add_table(doc,
        ["Attribute", "Value"],
        [
            ["Total Records", "10,194"],
            ["Original Columns", "18"],
            ["Unique Orders", "8,549"],
            ["Unique Customers", "5,044"],
            ["Unique Products", "15"],
            ["Cities Covered", "542"],
            ["States / Territories", "59"],
            ["Date Range", "January 1, 2025 -- September 9, 2025"],
            ["Order Years (Data)", "2024 -- 2025"],
            ["Country", "United States"],
        ])

    sub_heading(doc, "2.2  Feature Engineering (18 --> 43 Columns)")
    styled_para(doc,
        "A total of 25 new features were engineered from the raw data, transforming 18 columns into "
        "a comprehensive 43-column analytical dataset. The engineered features fall into the following "
        "categories:",
        10.5, color=TEXT, before=4, after=8)

    eng_categories = [
        ("Temporal Features (5):", "Order Year, Order Month, Order Quarter, Order Month Name, "
         "Shipping Days (calculated from Order Date to Ship Date)."),
        ("Route Classification Features (6):", "Route, Route_Region_State, Route_Factory_to_Region, "
         "Route_Factory_to_State, Route_to_Region, Route_to_State -- creating a multi-level "
         "route taxonomy from factory origin to customer destination."),
        ("Financial Ratio Features (5):", "Profit Margin (%), Cost Per Unit, Revenue Per Unit, "
         "Cost_Per_Ship_Day, Sales_Per_Ship_Day, Profit_Per_Ship_Day -- normalizing financial "
         "performance by volume and time dimensions."),
        ("Delay & Performance Features (5):", "Shipping_Lead_Time_Days, Lead_Time_Category, "
         "Delay_Bucket, Delay Flag (binary), Delay Status (categorical) -- classifying delivery "
         "performance against benchmarks."),
        ("Efficiency & Classification Features (4):", "Route Efficiency Score (composite 0-30 metric), "
         "Ship Category (Standard/Expedited), Shipping Lead Time (duplicate for analysis views)."),
    ]
    for title, body in eng_categories:
        bullet_block(doc, ">>", title, body, ACCENT)

    sub_heading(doc, "2.3  Route Efficiency Score Methodology")
    styled_para(doc,
        "The Route Efficiency Score is a composite metric (0-30 scale) that evaluates each shipment's "
        "overall routing performance. It incorporates multiple weighted factors including shipping lead "
        "time relative to the ship mode benchmark, on-time delivery status, cost efficiency ratio, "
        "and geographic complexity. A score of 30 represents optimal routing efficiency, while lower "
        "scores indicate opportunities for improvement. The dataset mean is 19.36 with a range of "
        "0.00 to 30.58.",
        10.5, color=TEXT, before=4, after=8)

    doc.add_page_break()

    # =========================================================
    # 3. METHODOLOGY & ANALYTICAL FRAMEWORK
    # =========================================================
    section_heading(doc, 3, "Methodology & Analytical Framework")

    sub_heading(doc, "3.1  Analytical Approach")
    styled_para(doc,
        "This project employs a multi-layered descriptive and diagnostic analytics approach, "
        "combining statistical aggregation, geographic analysis, temporal decomposition, and "
        "financial ratio analysis to deliver a 360-degree view of shipping route performance.",
        10.5, color=TEXT, before=4, after=8)

    approaches = [
        ("Descriptive Analytics:", "Statistical profiling of shipping times, delay rates, and "
         "efficiency scores across all dimensional cuts (region, state, ship mode, division, time)."),
        ("Geographic Analytics:", "Choropleth mapping and regional bottleneck analysis to identify "
         "spatial patterns in delivery performance across 59 U.S. states and territories."),
        ("Temporal Analytics:", "Monthly volume trending, seasonal heatmaps, quarterly efficiency "
         "comparison, and 30/90-day rolling average analysis to detect cyclical patterns."),
        ("Financial Analytics:", "Profit margin analysis, cost-per-ship-day evaluation, revenue "
         "efficiency measurement, and cost-speed frontier analysis across ship modes and divisions."),
        ("Diagnostic Analytics:", "Root cause analysis of delay patterns, outlier investigation, "
         "and drill-down capability from region to state to city to individual order level."),
    ]
    for title, body in approaches:
        bullet_block(doc, ">>", title, body, ACCENT)

    sub_heading(doc, "3.2  Technology Stack")
    add_table(doc,
        ["Component", "Technology", "Purpose"],
        [
            ["Data Processing", "Python, Pandas, NumPy", "Data cleaning, feature engineering, aggregation"],
            ["Visualization", "Plotly (interactive)", "Charts, maps, heatmaps, drill-downs"],
            ["Dashboard", "Streamlit", "Production-grade interactive web application"],
            ["BI Reporting", "Power BI (.pbix)", "Enterprise business intelligence dashboards"],
            ["Data Storage", "CSV, Excel (.xlsx)", "Processed dataset distribution"],
            ["Styling", "Custom CSS, Plotly Dark Theme", "Premium dark-mode UI/UX design"],
            ["Reporting", "python-docx, HTML/CSS", "Executive summary and project documentation"],
        ])

    doc.add_page_break()

    # =========================================================
    # 4. KEY PERFORMANCE INDICATORS (KPIs)
    # =========================================================
    section_heading(doc, 4, "Key Performance Indicators (KPIs)")

    styled_para(doc,
        "The following KPIs provide a high-level snapshot of Nassau Candy's shipping operations "
        "performance across the full dataset of 10,194 shipments.",
        10.5, color=TEXT, before=4, after=10)

    add_table(doc,
        ["KPI", "Value", "Benchmark / Context"],
        [
            ["Total Shipments", "10,194", "Across 59 states/territories, 542 cities"],
            ["Unique Orders", "8,549", "Avg 1.19 line items per order"],
            ["Avg Shipping Days", "1,320.84", "Extremely elevated (data includes 900+ day outliers)"],
            ["Median Shipping Days", "1,274", "All shipments exceed 900 days"],
            ["Std Dev Shipping Days", "262.44", "High variability in delivery times"],
            ["Min Shipping Days", "904", "Lowest recorded lead time"],
            ["Max Shipping Days", "1,642", "Highest recorded lead time"],
            ["On-Time Delivery Rate", "20.12%", "Critical: well below 85% industry benchmark"],
            ["Delayed Shipments", "8,143 (79.88%)", "Significant systemic delay pattern"],
            ["Avg Route Efficiency Score", "19.36 / 30", "Moderate efficiency with room for improvement"],
            ["Total Revenue", "$141,783.63", "Across all product lines"],
            ["Total Gross Profit", "$93,442.80", "66.51% average profit margin"],
            ["Total Cost", "$48,340.83", "Avg $1.24 cost per unit"],
            ["Unique Routes", "59", "Mapped region-state combinations"],
        ])

    doc.add_page_break()

    # =========================================================
    # 5. CRITICAL FINDINGS & INSIGHTS
    # =========================================================
    section_heading(doc, 5, "Critical Findings & Insights")

    findings = [
        ("CRITICAL", "Systemic Delivery Delays Across All Operations",
         "The on-time delivery rate stands at only 20.12%, meaning approximately 4 out of every 5 "
         "shipments are delayed. All 10,194 shipments have lead times exceeding 900 days, categorized "
         "as 'Very Slow (>2.5 Years)'. This represents a fundamental logistics challenge that spans "
         "all regions, ship modes, and product divisions. The delay distribution shows 2,051 orders in "
         "the 900-1000 day bucket, 4,764 in the 1200-1500 day bucket, and 3,379 in the 1500+ day bucket.",
         DANGER),
        ("CRITICAL", "Universal Lead Time Anomaly",
         "Every single shipment in the dataset has a shipping duration exceeding 900 days (minimum = 904, "
         "maximum = 1,642). This extreme uniformity suggests either a systemic measurement issue in the "
         "Ship Date recording process, a data pipeline defect, or an intentional pre-dating of orders "
         "during a legacy system migration. Regardless of root cause, this anomaly fundamentally distorts "
         "all time-based performance metrics and must be investigated as the highest priority action item.",
         DANGER),
        ("HIGH", "Regional Performance Disparities",
         "While all regions show high average shipping days, meaningful relative differences exist: "
         "Gulf region shows the best relative performance (1,311 avg days, 21.54% on-time, 19.94 efficiency) "
         "while Atlantic region underperforms (1,322 avg days, 18.75% on-time, 19.25 efficiency). The "
         "Interior region handles 2,335 orders with 19.96% on-time rate and the highest average profit "
         "margin (66.87%).",
         WARNING),
        ("HIGH", "Standard Class Dominance and Performance",
         "Standard Class shipping accounts for 6,120 of 10,194 shipments (60.04%), making it the "
         "overwhelmingly dominant ship mode. Ironically, Standard Class shows the best on-time rate "
         "(20.88%) and lowest average shipping days (1,314.33) among all modes. First Class (1,548 orders, "
         "18.93% on-time) and Same Day (547 orders, 16.82% on-time) actually perform worse on these "
         "metrics, suggesting the delay issue is systemic rather than mode-dependent.",
         WARNING),
        ("MEDIUM", "Concentrated Product Portfolio Risk",
         "The Chocolate division accounts for 9,844 of 10,194 orders (96.57%), creating extreme "
         "concentration risk. The remaining divisions -- Other (310 orders, 3.04%) and Sugar (40 orders, "
         "0.39%) -- are negligible in volume. Any disruption to the Chocolate supply chain would impact "
         "virtually the entire business.",
         BLUE),
        ("MEDIUM", "Geographic Volume Concentration",
         "California alone accounts for 2,001 orders (19.63%), followed by New York (1,128, 11.07%) "
         "and Texas (985, 9.66%). These top 3 states represent 40.36% of all shipments. The top 10 "
         "states account for approximately 70% of total volume, creating significant geographic "
         "concentration risk.",
         BLUE),
        ("INSIGHT", "Quarterly Volume Growth Pattern",
         "Order volume increases steadily across quarters: Q1 (1,415 orders), Q2 (2,153), Q3 (2,839), "
         "Q4 (3,787). Q4 volume is 2.68x Q1 volume, indicating strong seasonality likely driven by "
         "holiday demand for confectionery products. Despite the volume increase, delay rates remain "
         "relatively stable (79.27% - 81.28%), suggesting the delay issue is not capacity-driven.",
         SUCCESS),
    ]

    for severity, title, body, color in findings:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        r_badge = p.add_run(f"[{severity}]  ")
        r_badge.font.size = Pt(9); r_badge.font.bold = True
        r_badge.font.color.rgb = color; r_badge.font.name = FONT
        r_title = p.add_run(title)
        r_title.font.size = Pt(11); r_title.font.bold = True
        r_title.font.color.rgb = PRIMARY; r_title.font.name = FONT
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.left_indent = Cm(1.0)
        r_body = p2.add_run(body)
        r_body.font.size = Pt(10); r_body.font.name = FONT; r_body.font.color.rgb = TEXT_SEC

    doc.add_page_break()

    # =========================================================
    # 6. GEOGRAPHIC IMPACT ANALYSIS
    # =========================================================
    section_heading(doc, 6, "Geographic Impact Analysis")

    sub_heading(doc, "6.1  Regional Performance Comparison")
    add_table(doc,
        ["Region", "Orders", "% Share", "Avg Ship Days", "Avg Efficiency", "On-Time %", "Avg Margin %"],
        [
            ["Pacific",  "3,253", "31.91%", "1,322.19", "19.28", "20.78%", "66.59%"],
            ["Atlantic", "2,986", "29.29%", "1,322.75", "19.25", "18.75%", "66.20%"],
            ["Interior", "2,335", "22.91%", "1,323.09", "19.23", "19.96%", "66.87%"],
            ["Gulf",     "1,620", "15.89%", "1,311.37", "19.94", "21.54%", "66.42%"],
        ])

    sub_heading(doc, "6.2  Top 10 States by Shipment Volume")
    add_table(doc,
        ["State", "Orders", "Avg Ship Days", "Avg Efficiency Score"],
        [
            ["California",      "2,001", "1,318.43", "19.51"],
            ["New York",        "1,128", "1,324.49", "19.14"],
            ["Texas",             "985", "1,310.31", "20.01"],
            ["Pennsylvania",     "587", "1,324.20", "19.16"],
            ["Washington",       "506", "1,360.66", "16.93"],
            ["Illinois",         "492", "1,323.15", "19.22"],
            ["Ohio",             "469", "1,325.73", "19.06"],
            ["Florida",          "383", "1,311.40", "19.94"],
            ["Michigan",         "255", "1,300.56", "20.60"],
            ["North Carolina",   "249", "1,334.88", "18.51"],
        ])

    styled_para(doc,
        "Key Observation: Washington state stands out with the highest average shipping days (1,360.66) "
        "and lowest efficiency score (16.93) among top-10 states, suggesting a specific routing "
        "challenge for Pacific Northwest deliveries. Conversely, Michigan shows the best efficiency "
        "(20.60) and Texas shows the lowest average days (1,310.31) among high-volume states.",
        10.5, color=TEXT_SEC, before=4, after=8, indent=0.5)

    doc.add_page_break()

    # =========================================================
    # 7. SHIP MODE PERFORMANCE ANALYSIS
    # =========================================================
    section_heading(doc, 7, "Ship Mode Performance Analysis")

    add_table(doc,
        ["Ship Mode", "Orders", "% Share", "Avg Days", "Avg Efficiency", "On-Time %", "Avg Margin %"],
        [
            ["Standard Class", "6,120", "60.04%", "1,314.33", "19.35", "20.88%", "66.55%"],
            ["Second Class",   "1,979", "19.41%", "1,323.85", "19.37", "19.61%", "66.46%"],
            ["First Class",    "1,548", "15.19%", "1,338.28", "19.45", "18.93%", "66.46%"],
            ["Same Day",         "547",  "5.37%", "1,333.44", "19.26", "16.82%", "66.51%"],
        ])

    styled_para(doc,
        "Ship Category Split: Of the 10,194 shipments, 8,099 (79.45%) use Standard shipping and "
        "2,095 (20.55%) use Expedited shipping. Counter-intuitively, Standard Class outperforms all "
        "premium modes on both on-time rate and average shipping days. Same Day shipping shows the "
        "worst on-time rate (16.82%), suggesting the premium shipping infrastructure may not be "
        "delivering its intended speed advantage under the current systemic delay conditions.",
        10.5, color=TEXT_SEC, before=4, after=8, indent=0.5)

    styled_para(doc,
        "Profit margins are virtually identical across all ship modes (66.42% - 66.55%), indicating "
        "that shipping mode selection has minimal impact on order-level profitability in the current "
        "cost structure. This uniformity suggests potential opportunities for strategic mode upgrades "
        "without significant margin erosion.",
        10.5, color=TEXT_SEC, before=0, after=8, indent=0.5)

    # =========================================================
    # 8. PRODUCT & DIVISION ANALYSIS
    # =========================================================
    section_heading(doc, 8, "Product & Division Analysis")

    sub_heading(doc, "8.1  Division Distribution")
    add_table(doc,
        ["Division", "Orders", "% Share", "Avg Ship Days", "Avg Efficiency", "Avg Margin %"],
        [
            ["Chocolate", "9,844", "96.57%", "--", "--", "--"],
            ["Other",       "310",  "3.04%", "--", "--", "--"],
            ["Sugar",        "40",  "0.39%", "--", "--", "--"],
        ])

    sub_heading(doc, "8.2  Top Products by Volume")
    add_table(doc,
        ["Product Name", "Orders", "% of Total"],
        [
            ["Wonka Bar - Milk Chocolate",           "2,137", "20.96%"],
            ["Wonka Bar - Scrumdiddlyumptious",      "2,064", "20.25%"],
            ["Wonka Bar - Triple Dazzle Caramel",    "2,015", "19.77%"],
            ["Wonka Bar - Fudge Mallows",            "1,818", "17.83%"],
            ["Wonka Bar - Nutty Crunch Surprise",    "1,810", "17.76%"],
            ["Wonka Gum",                              "120",  "1.18%"],
            ["Kazookles",                               "96",  "0.94%"],
            ["Lickable Wallpaper",                      "94",  "0.92%"],
            ["Other products (7 items)",                "40",  "0.39%"],
        ])

    styled_para(doc,
        "The product portfolio is highly concentrated in five core Wonka Bar variants, which together "
        "constitute 96.57% of all orders. The remaining products (Wonka Gum, Kazookles, Lickable "
        "Wallpaper, Laffy Taffy, SweeTARTS, Fizzy Lifting Drinks, Nerds, Hair Toffee, Everlasting "
        "Gobstopper, Fun Dip) represent a long tail of low-volume specialty items.",
        10.5, color=TEXT_SEC, before=4, after=8, indent=0.5)

    doc.add_page_break()

    # =========================================================
    # 9. TEMPORAL & SEASONAL PATTERNS
    # =========================================================
    section_heading(doc, 9, "Temporal & Seasonal Patterns")

    sub_heading(doc, "9.1  Yearly Comparison")
    add_table(doc,
        ["Year", "Orders", "Avg Ship Days", "Avg Efficiency", "Total Sales"],
        [
            ["2024", "4,181", "1,094.02", "19.59", "$57,956.20"],
            ["2025", "6,013", "1,478.56", "19.20", "$83,827.43"],
        ])

    styled_para(doc,
        "Year-over-year, orders grew by 43.8% (4,181 to 6,013) while average shipping days increased "
        "by 35.2% (1,094 to 1,478 days). The efficiency score slightly decreased from 19.59 to 19.20, "
        "suggesting that the growth in volume may be straining the existing logistics infrastructure.",
        10.5, color=TEXT_SEC, before=4, after=8, indent=0.5)

    sub_heading(doc, "9.2  Quarterly Volume Distribution")
    add_table(doc,
        ["Quarter", "Orders", "Avg Ship Days", "Delay Rate"],
        [
            ["Q1", "1,415", "1,331.04", "79.58%"],
            ["Q2", "2,153", "1,323.14", "81.28%"],
            ["Q3", "2,839", "1,317.52", "79.78%"],
            ["Q4", "3,787", "1,318.21", "79.27%"],
        ])

    styled_para(doc,
        "Q4 accounts for 37.15% of annual volume, confirming strong holiday seasonality for "
        "confectionery products. Despite the 2.68x volume increase from Q1 to Q4, delay rates "
        "remain remarkably stable (79.27% - 81.28%), suggesting delays are structural rather "
        "than capacity-driven. Q2 shows the highest delay rate (81.28%).",
        10.5, color=TEXT_SEC, before=4, after=8, indent=0.5)

    # =========================================================
    # 10. FINANCIAL & PROFITABILITY ANALYSIS
    # =========================================================
    section_heading(doc, 10, "Financial & Profitability Analysis")

    add_table(doc,
        ["Financial Metric", "Value"],
        [
            ["Total Revenue (Sales)", "$141,783.63"],
            ["Total Gross Profit", "$93,442.80"],
            ["Total Cost (COGS)", "$48,340.83"],
            ["Average Profit Margin", "66.51%"],
            ["Average Revenue Per Unit", "$3.65"],
            ["Average Cost Per Unit", "$1.24"],
            ["Average Sales Per Order", "$13.91"],
            ["Total Units Shipped", "38,821"],
        ])

    styled_para(doc,
        "The business maintains a healthy 66.51% average gross profit margin across all operations. "
        "Profit margins are remarkably consistent across regions (66.20% - 66.87%) and ship modes "
        "(66.42% - 66.55%), indicating a stable pricing and cost structure. The average revenue per "
        "unit of $3.65 against an average cost of $1.24 provides a comfortable 2.94x markup ratio.",
        10.5, color=TEXT_SEC, before=4, after=8, indent=0.5)

    doc.add_page_break()

    # =========================================================
    # 11. RISK ASSESSMENT
    # =========================================================
    section_heading(doc, 11, "Risk Assessment")

    risks = [
        ("CRITICAL", "Data Integrity & Measurement Risk",
         "All 10,194 shipments show 900+ day lead times, suggesting a systemic data recording issue. "
         "If Ship Date data is inaccurate, all time-based KPIs (shipping days, delay rates, efficiency "
         "scores) may be unreliable. Immediate audit of the Ship Date recording process is essential "
         "before making operational decisions based on these metrics."),
        ("HIGH", "Single-Factory Distribution Vulnerability",
         "All shipments originate from a single northeastern factory, creating a single point of failure. "
         "Any disruption (weather, labor, infrastructure) would halt all 10,194+ shipments. The average "
         "1,320-day lead time amplifies this risk by extending the recovery timeline."),
        ("HIGH", "Product Concentration Risk",
         "With 96.57% of orders in the Chocolate division and 96.57% in just 5 Wonka Bar products, "
         "any supply chain disruption, ingredient shortage, or demand shift would impact virtually "
         "the entire business."),
        ("HIGH", "Systemic Delay Rate (79.88%)",
         "Nearly 4 out of 5 shipments are delayed regardless of ship mode, region, or product. "
         "This pattern suggests root causes beyond logistics optimization -- potentially measurement "
         "methodology, benchmark definition, or fulfillment process issues."),
        ("MEDIUM", "Geographic Concentration",
         "Top 3 states (CA, NY, TX) account for 40.36% of all orders. Economic downturns or "
         "regulatory changes in these states would disproportionately impact revenue."),
        ("MEDIUM", "Seasonal Capacity Planning",
         "Q4 volume is 2.68x Q1, yet no evidence of proactive capacity planning exists. Holiday "
         "season surges could overwhelm existing logistics if not managed proactively."),
        ("LOW", "Ship Mode Cost Escalation",
         "Upgrading ship modes for efficiency gains may increase per-order costs by 5-8%, though "
         "current data shows margins are uniform across modes (66.42% - 66.55%)."),
    ]

    for level, title, body in risks:
        colors = {"CRITICAL": DANGER, "HIGH": WARNING, "MEDIUM": BLUE, "LOW": SUCCESS}
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        r_b = p.add_run(f"[{level}]  ")
        r_b.font.size = Pt(10); r_b.font.bold = True
        r_b.font.color.rgb = colors.get(level, TEXT); r_b.font.name = FONT
        r_t = p.add_run(title)
        r_t.font.size = Pt(11); r_t.font.bold = True
        r_t.font.color.rgb = PRIMARY; r_t.font.name = FONT
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.left_indent = Cm(1.0)
        r_d = p2.add_run(body)
        r_d.font.size = Pt(10); r_d.font.name = FONT; r_d.font.color.rgb = TEXT_SEC

    doc.add_page_break()

    # =========================================================
    # 12. STRATEGIC RECOMMENDATIONS
    # =========================================================
    section_heading(doc, 12, "Strategic Recommendations")

    recs = [
        ("P0", "HIGH", "Data Quality Audit & Ship Date Investigation",
         "Immediately investigate the Ship Date recording methodology to determine why all shipments "
         "show 900+ day lead times. Audit the data pipeline from order entry to shipment tracking. "
         "Validate against carrier tracking data. This is the foundational prerequisite for all other "
         "optimization efforts -- without accurate lead time data, no meaningful performance improvement "
         "can be measured.",
         "Foundational -- enables all other recommendations"),
        ("P1", "HIGH", "Regional Warehouse Expansion",
         "Establish a secondary distribution center in the Interior or Gulf region (Dallas-Fort Worth "
         "or Kansas City) to reduce average shipping distances by 30-40% for underserved corridors. "
         "This would create redundancy against single-factory risk and improve service levels for "
         "2,335 Interior region orders.",
         "30-40% distance reduction for Interior region"),
        ("P1", "HIGH", "Dynamic Ship Mode Allocation",
         "Implement rule-based automatic upgrades from Standard to Second Class for high-delay-probability "
         "routes and high-value orders. Current data shows margins are uniform across modes (66.42% - "
         "66.55%), meaning upgrades can be absorbed with minimal margin impact.",
         "Improved on-time rates with <1% margin impact"),
        ("P2", "MEDIUM", "SLA-Based Routing Protocols",
         "Define explicit Service Level Agreements for each Region x Ship Mode combination with "
         "automated alerts when 30-day rolling performance drops below targets. Deploy real-time "
         "monitoring through the Streamlit dashboard.",
         "Proactive performance management"),
        ("P2", "MEDIUM", "Seasonal Capacity Pre-Planning",
         "Pre-secure additional carrier contracts and pre-position Chocolate division inventory 60 "
         "days before Q4 holiday season. Q4 represents 37.15% of annual volume (3,787 orders) -- "
         "proactive planning is essential to maintain service levels during peak demand.",
         "Prevents Q4 capacity shortfalls"),
        ("P2", "MEDIUM", "Product Diversification Strategy",
         "The 96.57% concentration in Chocolate creates extreme risk. Evaluate expansion of Sugar "
         "and Other division product lines to build a more resilient revenue base. Current Sugar "
         "division (40 orders) and Other (310 orders) represent untapped growth potential.",
         "Reduced concentration risk"),
        ("P3", "LOW", "Premium Shipping for High-Value SKUs",
         "Default Chocolate division orders to First Class for destinations beyond 500 miles. "
         "With an average revenue per unit of $3.65 and cost per unit of $1.24, the 2.94x markup "
         "provides margin cushion to absorb incremental shipping costs.",
         "Improved customer experience for high-value orders"),
    ]

    for priority, impact, title, body, expected in recs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        impact_colors = {"HIGH": SUCCESS, "MEDIUM": WARNING, "LOW": BLUE}
        r_p = p.add_run(f"[{priority}] ")
        r_p.font.size = Pt(9); r_p.font.bold = True
        r_p.font.color.rgb = DANGER if priority == "P0" else ACCENT; r_p.font.name = FONT
        r_t = p.add_run(title)
        r_t.font.size = Pt(11); r_t.font.bold = True
        r_t.font.color.rgb = PRIMARY; r_t.font.name = FONT
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.left_indent = Cm(1.0)
        r_d = p2.add_run(body)
        r_d.font.size = Pt(10); r_d.font.name = FONT; r_d.font.color.rgb = TEXT_SEC
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(10)
        p3.paragraph_format.left_indent = Cm(1.0)
        r_e = p3.add_run(f"Expected Impact: {expected}")
        r_e.font.size = Pt(9); r_e.font.bold = True
        r_e.font.color.rgb = impact_colors.get(impact, TEXT); r_e.font.name = FONT

    doc.add_page_break()

    # =========================================================
    # 13. IMPLEMENTATION ROADMAP
    # =========================================================
    section_heading(doc, 13, "Implementation Roadmap")

    phases = [
        ("PHASE 1  --  Immediate (0-30 Days)", [
            ("Data Quality Audit", "Investigate Ship Date recording anomalies; validate against carrier systems."),
            ("SLA Definition", "Establish Region x Ship Mode performance targets and monitoring thresholds."),
            ("Dashboard Deployment", "Deploy Streamlit dashboard for real-time KPI tracking by stakeholders."),
            ("Outlier Investigation", "Forensic analysis of 10,194 extreme lead-time records."),
        ]),
        ("PHASE 2  --  Short-Term (30-90 Days)", [
            ("Dynamic Ship Mode Pilot", "Launch auto-upgrade pilot on top 10 highest-delay routes."),
            ("Monitoring Infrastructure", "Implement automated alerts for SLA breaches and trend shifts."),
            ("Carrier Performance Review", "Audit carrier performance data to validate internal metrics."),
        ]),
        ("PHASE 3  --  Medium-Term (90-180 Days)", [
            ("Seasonal Capacity Planning", "Pre-secure Q4 carrier contracts; pre-position holiday inventory."),
            ("Product Diversification Analysis", "Evaluate Sugar/Other division expansion opportunities."),
            ("Route Optimization Modeling", "Build predictive models for route-level lead time estimation."),
        ]),
        ("PHASE 4  --  Long-Term (6-18 Months)", [
            ("Regional Warehouse Feasibility", "Conduct site analysis for Interior region distribution center."),
            ("Warehouse Deployment", "If approved, proceed with facility procurement and setup."),
            ("End-to-End Supply Chain Digitization", "Integrate order, shipping, and carrier data into unified platform."),
        ]),
    ]

    for phase_title, items in phases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(phase_title)
        r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = ACCENT; r.font.name = FONT
        for item_title, item_desc in items:
            pi = doc.add_paragraph()
            pi.paragraph_format.space_before = Pt(2)
            pi.paragraph_format.space_after = Pt(2)
            pi.paragraph_format.left_indent = Cm(1.0)
            ri_t = pi.add_run(f">> {item_title}: ")
            ri_t.font.size = Pt(10); ri_t.font.bold = True
            ri_t.font.color.rgb = PRIMARY; ri_t.font.name = FONT
            ri_d = pi.add_run(item_desc)
            ri_d.font.size = Pt(10); ri_d.font.color.rgb = TEXT_SEC; ri_d.font.name = FONT

    doc.add_paragraph()

    # =========================================================
    # 14. TECHNICAL ARCHITECTURE & DELIVERABLES
    # =========================================================
    section_heading(doc, 14, "Technical Architecture & Deliverables")

    sub_heading(doc, "14.1  Interactive Dashboard (Streamlit)")
    styled_para(doc,
        "A production-grade interactive web dashboard built with Streamlit and Plotly, featuring "
        "a premium dark-mode UI design with glassmorphism effects, gradient KPI cards, and smooth "
        "hover animations. The dashboard provides 6 analytical tabs:",
        10.5, color=TEXT, before=4, after=8)

    tabs = [
        ("Tab 1 - Route Efficiency Overview:", "Top 20 routes by lead time (horizontal bar chart) + "
         "Route Performance Leaderboard (sortable, color-coded data table)."),
        ("Tab 2 - Geographic Shipping Map:", "Dual choropleth maps (Avg Shipping Days + Efficiency Score) "
         "covering all 59 states/territories + Regional Bottleneck Analysis with bar charts."),
        ("Tab 3 - Ship Mode Comparison:", "KPI row per mode + grouped bar charts + box plot distributions "
         "+ sunburst chart (Ship Mode to Delay Status) + summary data table."),
        ("Tab 4 - Route Drill-Down:", "State-level deep dive with 6 KPIs + top cities bar chart + "
         "shipment timeline scatter plot + order-level detail table with progress bars."),
        ("Tab 5 - Trend & Seasonality:", "Monthly volume/lead-time dual-axis chart + Ship Mode x Month "
         "delay heatmap + quarterly efficiency comparison + 30/90-day rolling averages."),
        ("Tab 6 - Profitability & Cost:", "Profit margin violin plots (by region/mode) + cost efficiency "
         "frontier scatter + revenue/cost per unit by division + cost/profit per ship day analysis."),
    ]
    for title, body in tabs:
        bullet_block(doc, ">>", title, body, ACCENT)

    sub_heading(doc, "14.2  Sidebar Filters")
    styled_para(doc,
        "The dashboard features a comprehensive sidebar with filters for: Date Range (date picker), "
        "Region & State (multi-select), Ship Mode (multi-select), Division (multi-select), and "
        "Lead-Time Threshold (range slider). An Executive Summary download button is also included.",
        10.5, color=TEXT, before=4, after=8)

    sub_heading(doc, "14.3  Project Deliverables")
    add_table(doc,
        ["Deliverable", "Format", "Description"],
        [
            ["Interactive Dashboard", "Streamlit (app.py)", "6-tab analytics dashboard with real-time filters"],
            ["Preprocessed Dataset", "CSV + XLSX", "43-column engineered dataset (10,194 records)"],
            ["Power BI Report", ".pbix", "Enterprise BI dashboard for executive consumption"],
            ["Executive Summary", "HTML + DOCX", "Formatted report for stakeholder review"],
            ["Project Summary", "DOCX", "Comprehensive in-depth project documentation"],
            ["Source Data", "CSV", "Original 18-column raw dataset"],
        ])

    doc.add_page_break()

    # =========================================================
    # 15. CONCLUSION
    # =========================================================
    section_heading(doc, 15, "Conclusion")

    styled_para(doc,
        "This comprehensive analysis of Nassau Candy Distributor's factory-to-customer shipping routes "
        "reveals both significant challenges and clear opportunities for optimization across a portfolio "
        "of 10,194 shipments spanning 59 U.S. states and territories.",
        11, color=TEXT, before=6, after=10)

    styled_para(doc,
        "The most critical finding is the systemic nature of the delivery delay pattern -- with a "
        "79.88% delay rate and all shipments showing 900+ day lead times -- which strongly suggests "
        "a data quality or measurement methodology issue that must be resolved before any operational "
        "improvements can be accurately measured. This data integrity investigation should be the "
        "immediate, highest-priority action.",
        10.5, color=TEXT, before=0, after=10)

    styled_para(doc,
        "Beyond the data quality concern, the analysis reveals meaningful relative performance "
        "differences across regions (Gulf outperforming Atlantic), counter-intuitive ship mode "
        "dynamics (Standard Class outperforming premium modes), extreme product concentration "
        "(96.57% Chocolate), and strong seasonal patterns (Q4 = 2.68x Q1 volume). These insights "
        "provide a foundation for targeted logistics interventions once the underlying data quality "
        "is validated.",
        10.5, color=TEXT, before=0, after=10)

    styled_para(doc,
        "The project delivers a complete analytical toolkit -- an interactive Streamlit dashboard "
        "with 6 analytical tabs, a 43-feature engineered dataset, Power BI reports, and actionable "
        "executive documentation -- enabling Nassau Candy's stakeholders to make data-driven "
        "decisions about warehouse expansion, ship mode optimization, seasonal planning, and "
        "product diversification.",
        10.5, color=TEXT, before=0, after=10)

    styled_para(doc,
        "With the recommended phased implementation roadmap (0-18 months), Nassau Candy can "
        "systematically address its logistics challenges, starting with data quality remediation "
        "and progressing through route optimization, capacity planning, and infrastructure expansion "
        "to achieve sustainable improvements in delivery performance and customer satisfaction.",
        10.5, color=TEXT, before=0, after=12)

    # =========================================================
    # FOOTER
    # =========================================================
    divider(doc)

    styled_para(doc,
        "Nassau Candy Distributor -- Factory-to-Customer Shipping Route Efficiency Analysis",
        9, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=2)
    styled_para(doc,
        "Project Summary Report  |  Prepared by Om Ahir  |  Data Analytics Division  |  July 2026",
        9, color=TEXT_SEC, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
    styled_para(doc,
        "This document is confidential and intended for authorized government stakeholder review only. "
        "Distribution outside authorized channels is prohibited.",
        8, italic=True, color=TEXT_SEC, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    # ── Save ──
    out_dir  = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "Project_Summary_Nassau_Candy.docx")
    doc.save(out_path)
    print(f"[OK] Project Summary generated successfully!")
    print(f"   File: {out_path}")
    print(f"   Format: Google Docs-compatible .docx")
    print(f"   Sections: 15 sections, ~20 pages")
    return out_path


if __name__ == "__main__":
    generate()
